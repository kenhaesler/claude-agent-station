"""Text extraction for non-native vision attachment types (spec 2026-05-21)."""
import io

import pytest
from openpyxl import Workbook

from app.services.vision_attachments import extract_text, EXTRACTION_MAX_BYTES


def _xlsx_bytes() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "qty", "price"])
    ws.append(["apple", 3, 0.5])
    ws.append(["banana", 5, 0.25])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _csv_bytes() -> bytes:
    return b"name,qty\napple,3\nbanana,5\n"


def _docx_bytes() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_xlsx_renders_markdown_table():
    text = extract_text(
        _xlsx_bytes(),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert "Sheet1" in text
    assert "| name | qty | price |" in text
    assert "| apple | 3 | 0.5 |" in text


def test_extract_csv_passes_through():
    text = extract_text(_csv_bytes(), mime="text/csv")
    assert "apple,3" in text
    assert "banana,5" in text


def test_extract_docx_paragraphs():
    text = extract_text(_docx_bytes(), mime=(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ))
    assert "Hello world." in text
    assert "Second paragraph." in text


def test_extract_returns_none_for_native_types():
    assert extract_text(b"\x89PNG\r\n", mime="image/png") is None
    assert extract_text(b"%PDF-1.4", mime="application/pdf") is None


def test_extract_truncates_large_output():
    wb = Workbook()
    ws = wb.active
    ws.append(["col"])
    for i in range(50_000):
        ws.append([f"row-{i}-padding-padding-padding"])
    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text(
        buf.getvalue(),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert text is not None
    assert len(text) <= EXTRACTION_MAX_BYTES + 200
    assert "[truncated" in text
